import logging
import os
import re
from abc import ABC, abstractmethod
from pathlib import Path

import PyPDF2
import openpyxl
import pandas as pd
from docx import Document

logger = logging.getLogger(__name__)


class TextExtractor(ABC):
    """Abstract base class for text extractors"""

    @abstractmethod
    def extract(self, file_path: str) -> str:
        pass

    @abstractmethod
    def get_file_type(self) -> str:
        pass


class PDFExtractor(TextExtractor):
    """Extract text from PDF files"""

    def extract(self, file_path: str) -> str:
        try:
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise

    def get_file_type(self) -> str:
        return "PDF"


class DOCXExtractor(TextExtractor):
    """Extract text from DOCX files"""

    def extract(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"

            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise

    def get_file_type(self) -> str:
        return "DOCX"


class XLSXExtractor(TextExtractor):
    """Extract text from XLSX files"""

    def extract(self, file_path: str) -> str:
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"

                data = []
                for row in sheet.iter_rows(values_only=True):
                    data.append(row)

                if data:
                    df = pd.DataFrame(data)
                    text += df.to_string(index=False, na_rep="") + "\n"

            return text
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            raise

    def get_file_type(self) -> str:
        return "XLSX"


class TXTExtractor(TextExtractor):
    """Extract text from TXT files"""

    def extract(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            encodings = ["latin-1", "cp1252", "iso-8859-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise Exception("Unable to decode text file with common encodings")
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            raise

    def get_file_type(self) -> str:
        return "TXT"


class TextCleaner:
    """Clean and preprocess extracted text"""

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean text by removing artifacts and normalizing content"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove common PDF artifacts
        text = re.sub(r"[^\x00-\x7F]+", " ", text)  # Remove non-ASCII
        text = re.sub(r"\x0c", "", text)  # Remove form feed

        # Remove excessive punctuation
        text = re.sub(r"[.]{3,}", "...", text)
        text = re.sub(r"[-]{3,}", "---", text)

        # Remove duplicate lines (common in extracted content)
        lines = text.split("\n")
        seen = set()
        unique_lines = []
        for line in lines:
            line_stripped = line.strip()
            if line_stripped and line_stripped not in seen:
                seen.add(line_stripped)
                unique_lines.append(line)

        # Join and final cleanup
        text = "\n".join(unique_lines)
        text = text.strip()

        return text

    @staticmethod
    def count_words(text: str) -> int:
        """Count words in text"""
        if not text:
            return 0
        return len(text.split())


class DocumentParser:

    def __init__(self):
        self.extractors = {
            ".pdf": PDFExtractor(),
            ".docx": DOCXExtractor(),
            ".xlsx": XLSXExtractor(),
            ".xls": XLSXExtractor(),
            ".txt": TXTExtractor(),
        }
        self.cleaner = TextCleaner()
        self.document_text: str = ""

    def process_document(
        self,
        file_path=None,
        raw_text=None,
    ):
        """Process a document and return structured summary"""
        if raw_text:
            self.document_text = raw_text.strip()

            # Clean text
            # self.document_text = self.cleaner.clean_text(raw_text)
            word_count = self.cleaner.count_words(self.document_text)

            logger.info(f"Extracted {word_count} words ")
        elif file_path:
            collected_text = ""
            if isinstance(file_path, str):
                file_paths = [file_path]
            else:
                file_paths = file_path
            for file_path in file_paths:
                # Validate file exists
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"File not found: {file_path}")

                # Extract file extension
                file_extension = Path(file_path).suffix.lower()

                if file_extension not in self.extractors:
                    raise ValueError(f"Unsupported file type: {file_extension}")

                extractor = self.extractors[file_extension]

                logger.info(f"Processing {file_extension} file: {file_path}")

                # Extract text
                raw_text = extractor.extract(file_path)
                # Clean text
                # cleaner_text = self.cleaner.clean_text(raw_text)

                collected_text += raw_text + "\n"
                word_count = self.cleaner.count_words(collected_text)

                logger.info(
                    f"Extracted {word_count} words from {extractor.get_file_type()}"
                )

                # TODO:- We can handle chunking later if required.
            self.document_text = collected_text
